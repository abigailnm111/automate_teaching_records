
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

import openpyxl

import re
import os

from hidden_variables import faculty_names, file_name


class evaluationScores:
    def __init__(self, faculty):
        self.faculty= faculty
        self.all_scores={}
        self.search_name=re.search("\S+,\s\S", self.faculty).group() #gets last name and first initial with comma and space
        
    def save_scores(self, rundown, q):
        self.courses=[]
        name_column=get_quarter_columns(rundown,"Instructor Name")
        upper_name=self.search_name.upper() 
        self.dept_ins_avg=rundown[get_quarter_columns(rundown, "Dept Inst AVG")+"2"].value
        self.dept_crs_avg=rundown[get_quarter_columns(rundown, "Dept Crs AVG")+"2"].value
        for row in rundown[name_column]:
            r= str(row.row)
            if upper_name in row.value:
                self.course_id= rundown[get_quarter_columns(rundown, "Subject Course Section")+r].value[:-6]
                self.title= rundown[get_quarter_columns(rundown, "Course Title")+r].value
                self.enrollment= rundown[get_quarter_columns(rundown, "Enrollment")+r].value
                self.response= rundown[get_quarter_columns(rundown, "Response Rate")+r].value
                self.ins_avg= rundown[get_quarter_columns(rundown, "Inst AVG")+r].value
                self.crs_avg=rundown[get_quarter_columns(rundown, "Crs AVG")+r].value
                self.courses.append([self.course_id, self.title, self.enrollment, self.response, self.ins_avg, 
                                     self.crs_avg, self.dept_ins_avg, self.dept_crs_avg])
            self.all_scores[q]= self.courses
        return self.all_scores
        
def get_quarter_columns(rundown, column_name):
        for column in rundown[1]:
            if re.search(column_name, column.value)!= None:
                return column.column_letter
            
def get_quarters_years():
    years= ['19','20','21']
    sessions=["W", "S", "F"]
    quarter_list=[]
    for y in years:
        for q in sessions:
            quarter_list.append(y+q)
    return quarter_list
 
def open_rundown_file(yq):
    path=os.path.join("ENGL Evaluations", yq+" ENGL/Rundown Reports")
    rundown_file=os.path.join(path, yq+file_name)
    if os.path.isfile(rundown_file)== True:
        rundown_report= openpyxl.load_workbook(rundown_file)
        rundown= rundown_report.worksheets[0]
        return rundown
    else:
        return None


def write_teaching_record(faculty):
     name=faculty.faculty
     template=Document('Teaching Record.docx')
     header_section=template.sections[-1]
     #doc= Document('new temp.docx')
     for paragraph in header_section.header.paragraphs:
         new_paragraph= re.sub(r"<NAME>", name, paragraph.text)
         paragraph.text=new_paragraph
     score_table=template.tables[0]
     cell=score_table.cell(0,7)
     xml=cell._tc.xml
     i=0
     for quarter in faculty.all_scores:
            row_cells=score_table.rows[i].cells
            row_cells[0].text=quarter
            for course in faculty.all_scores[quarter]:
                row_cells=score_table.rows[i].cells
                
                row_cells[1].text=course[0] #course_id
                row_cells[2].text= course[1] #title
                row_cells[4].text= str(course[2]) #enrollment
                row_cells[6].text= "{:.2f}".format(course[4]) #inst avg
                row_cells[7].text= "{:.2f}".format(course[5]) #crs avg
                row_cells[8].text= "{:.2%}".format(course[3]) # response rate
                row_cells[9].text= "{:.2f}".format(course[6])# dept inst av
                row_cells[10].text= "{:.2f}".format(course[7])# dept crs avg
                i+=1
                score_table.add_row()
            if faculty.all_scores[quarter]==[]:
                i+=1
                score_table.add_row()
            i+=1
            score_table.add_row()
     for row in score_table.rows:
        row_cells=row.cells
        for x in range(6,11):
            shade=re.search('w:fill.+"',xml).group()
            shade=parse_xml(r'<w:shd {} {}/>'.format(nsdecls('w'),shade))
            row_cells[x]._tc.get_or_add_tcPr().append(shade)
     template.save(faculty.faculty+"_Teaching Record.docx")

def main():
    quarters= get_quarters_years()
    #iterate through each faculty member
    all_fac=[]
    for name in faculty_names:
            faculty=evaluationScores(name)   
            all_fac.append(faculty)
    for name in all_fac:  
        #iterate through each quarter
        for q in quarters:
            #open rundown file per quarter
            rundown=open_rundown_file(q)
            if rundown!=None:
                name.save_scores(rundown,q)
        write_teaching_record(name)
        
        
if __name__ == '__main__':
    main()
      

    
   
    


